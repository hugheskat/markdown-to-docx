from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.xmlchemy import OxmlElement
import os
import time
import endmessage
import scripttime

# get the start time
start_time  = time.time()

# log usage
logusage.addtolog(getpass.getuser(), time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()), os.path.basename(__file__))

## regex patterns to capture formatting, inline code, hyperlinks and numbered lists
inline_code_pattern = re.compile(r"([`].*?[`])")
inline_code_pattern2 = re.compile(r"[`](.*?)[`]")
bold_text_pattern = re.compile(r'(\*{2}.*?\*{2})')
bold_text_pattern2 = re.compile(r'\*{2}(.*?)\*{2}')
italic_text_pattern = re.compile(r'(\*{1}.*?\*{1})')
italic_text_pattern2 = re.compile(r'\*{1}(.*?)\*{1}')
italic_bold_text_pattern = re.compile(r'(\*{3}.*?\*{3})')
italic_bold_text_pattern2 = re.compile(r'\*{3}(.*?)\*{3}')
# hyperlink..
hyperlink_name = "[^]]+" # anything that isn't a square closing bracket
hyperlink_url= "http[s]?://[^)]+" # http:// or https:// followed by anything but a closing parenthesis
hyperlink_pattern = re.compile(r'(\[{0}\]\(\s*{1}\s*\))'.format(hyperlink_name, hyperlink_url))
hyperlink_url_pattern = re.compile(r"(http[s]?://[^)]+)")
hyperlink_name_pattern = re.compile(r'([^]]+)')
# ordered lists
level2_ordered_list = re.compile(r'(\s{2}[0-9a-zA-Z]{,3}[.])')
level3_ordered_list = re.compile(r'(\s{4}[0-9a-zA-Z]{,3}[.])')

def add_hyperlink(paragraph, url, text, color, underline):
    import docx

    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    
    This function was taken from: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410.
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Add underlining
    if underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'single')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def format_para(line, paragraph):
    """This converts any formatting/inline code/hyperlinks within a paragraph object
    (eg. within list items and block quotes etc).

    : param line: The markdown line.
    : param paragraph: The docx paragraph object.
    """
    # convert hyperlinks
    if re.findall(hyperlink_pattern, line):
        hyperlink = re.split(hyperlink_pattern, line)
        for text_ in hyperlink:
            if text_ is not None:
                if re.match(hyperlink_pattern, text_):
                    # get the url
                    url_ = re.split(hyperlink_url_pattern, text_)[-2]
                    # get the hyperlink text
                    name_ = ""
                    if re.split(hyperlink_name_pattern, text_)[0] != '' and \
                            re.split(hyperlink_name_pattern, text_)[0].startswith('['):
                        name_ = re.split(hyperlink_name_pattern, text_)[0].strip('[')
                    elif re.split(hyperlink_name_pattern, text_)[1] != '' and \
                            re.split(hyperlink_name_pattern, text_)[1].startswith('['):
                        name_ = re.split(hyperlink_name_pattern, text_)[1].strip('[')
                    # add in the hyperlink with the text blue and underlined
                    add_hyperlink(paragraph, url_, name_, '334BFF', True)
                # if bold, italic text is in the line, format bold, italic
                elif re.findall(italic_bold_text_pattern, text_):
                    italic_bold_text = re.split(italic_bold_text_pattern, text_)
                    for text2 in italic_bold_text:
                        if text2 is not None:
                            if re.match(italic_bold_text_pattern2, text2):
                                text2 = text2.replace("***", "")
                                # convert to bold italic and add to paragraph object
                                para_run = paragraph.add_run(text2)
                                para_run.italic = True
                                para_run.bold = True
                            # or if bold text is in this section of text, format bold
                            elif re.findall(bold_text_pattern, text2):
                                bold_text_split = re.split(bold_text_pattern, text2)
                                for text3 in bold_text_split:
                                    if text3 is not None:
                                        if re.match(bold_text_pattern2, text3):
                                            # convert to bold and add to paragraph object
                                            text3 = text3.replace("**", "")
                                            para_run = paragraph.add_run(text3)
                                            para_run.bold = True
                                        # or if italic text is in this section of text, format italic
                                        elif re.findall(italic_text_pattern, text3):
                                            italic_text_split = re.split(italic_text_pattern, text3)
                                            for text4 in italic_text_split:
                                                if text4 is not None:
                                                    if re.match(italic_text_pattern2, text4):
                                                        # convert to italic and add to paragraph object
                                                        text4 = text4.replace("*", "")
                                                        para_run = paragraph.add_run(text4)
                                                        para_run.italic = True
                                                    # or if inline code is in this section of text, format inline code
                                                    elif re.findall(inline_code_pattern, text4):
                                                        inline_code_split = re.split(inline_code_pattern, text4)
                                                        for text5 in inline_code_split:
                                                            if text5 is not None:
                                                                if re.match(inline_code_pattern2, text5):
                                                                    # change formatting and add to paragraph object
                                                                    text5 = text5.replace("`", "")
                                                                    para_run = paragraph.add_run(text5)
                                                                    font = para_run.font
                                                                    font.name = 'consolas'
                                                                    font.size = Pt(9)
                                                                    font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                                else:
                                                                    paragraph.add_run(text5)
                                                    else:
                                                        paragraph.add_run(text4)
                                        # or if inline code is in this section of text, format inline code
                                        elif re.findall(inline_code_pattern, text3):
                                            inline_code_split = re.split(inline_code_pattern, text3)
                                            for text4 in inline_code_split:
                                                if text4 is not None:
                                                    if re.match(inline_code_pattern2, text4):
                                                        # change formatting and add to paragraph object
                                                        text4 = text4.replace("`", "")
                                                        para_run = paragraph.add_run(text4)
                                                        font = para_run.font
                                                        font.name = 'consolas'
                                                        font.size = Pt(9)
                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                    else:
                                                        paragraph.add_run(text4)
                                        else:
                                            paragraph.add_run(text3)
                            # or if italic text is in this section of text, format italic
                            elif re.findall(italic_text_pattern, text2):
                                italic_text_split = re.split(italic_text_pattern, text2)
                                for text3 in italic_text_split:
                                    if text3 is not None:
                                        if re.match(italic_text_pattern2, text3):
                                            # convert to italic and add to paragraph object
                                            text3 = text3.replace("*", "")
                                            para_run = paragraph.add_run(text3)
                                            para_run.italic = True
                                        # or if inline code is in this section of text, format inline code
                                        elif re.findall(inline_code_pattern, text3):
                                            inline_code_split = re.split(inline_code_pattern, text3)
                                            for text4 in inline_code_split:
                                                if text4 is not None:
                                                    if re.match(inline_code_pattern2, text4):
                                                        # change formatting and add to paragraph object
                                                        text4 = text4.replace("`", "")
                                                        para_run = paragraph.add_run(text4)
                                                        font = para_run.font
                                                        font.name = 'consolas'
                                                        font.size = Pt(9)
                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                    else:
                                                        paragraph.add_run(text4)
                                        else:
                                            paragraph.add_run(text3)
                            # or if inline code is in this section of text, format inline code
                            elif re.findall(inline_code_pattern, text2):
                                inline_code_split = re.split(inline_code_pattern, text2)
                                for text3 in inline_code_split:
                                    if text3 is not None:
                                        if re.match(inline_code_pattern2, text3):
                                            # change formatting and add to paragraph object
                                            text3 = text3.replace("`", "")
                                            para_run = paragraph.add_run(text3)
                                            font = para_run.font
                                            font.name = 'consolas'
                                            font.size = Pt(9)
                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # or if bold text is in the line, format bold
                elif re.findall(bold_text_pattern, text_):
                    # add para to docx document
                    bold_text_split = re.split(bold_text_pattern, text_)
                    for text2 in bold_text_split:
                        if text2 is not None:
                            if re.match(bold_text_pattern2, text2):
                                # format bold and add to paragraph object
                                text2 = text2.replace("**", "")
                                para_run = paragraph.add_run(text2)
                                para_run.bold = True
                            # or if italic text is in this section of text, format italic
                            elif re.findall(italic_text_pattern, text2):
                                italic_text_split = re.split(italic_text_pattern, text2)
                                for text3 in italic_text_split:
                                    if text3 is not None:
                                        if re.match(italic_text_pattern2, text3):
                                            # format italic and add to paragraph object
                                            text3 = text3.replace("*", "")
                                            para_run = paragraph.add_run(text3)
                                            para_run.italic = True
                                        # if inline code is in this section of text, format italic
                                        elif re.findall(inline_code_pattern, text3):
                                            inline_code_split = re.split(inline_code_pattern, text3)
                                            for text4 in inline_code_split:
                                                if text4 is not None:
                                                    if re.match(inline_code_pattern2, text4):
                                                        # change formatting and add to paragraph object
                                                        text4 = text4.replace("`", "")
                                                        para_run = paragraph.add_run(text4)
                                                        font = para_run.font
                                                        font.name = 'consolas'
                                                        font.size = Pt(9)
                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                    else:
                                                        paragraph.add_run(text4)
                                        else:
                                            paragraph.add_run(text3)
                            # or if inline code is in this section of text, format inline code
                            elif re.findall(inline_code_pattern, text2):
                                inline_code_split = re.split(inline_code_pattern, text2)
                                for text3 in inline_code_split:
                                    if text3 is not None:
                                        if re.match(inline_code_pattern2, text3):
                                            # change formatting and add to paragraph object
                                            text3 = text3.replace("`", "")
                                            para_run = paragraph.add_run(text3)
                                            font = para_run.font
                                            font.name = 'consolas'
                                            font.size = Pt(9)
                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # inline code
                # change the font of inline code in the line to Consolas and highlight in light grey
                elif re.findall(inline_code_pattern, text_):
                    inline_code_split = re.split(inline_code_pattern, text_)
                    for text2 in inline_code_split:
                        if text2 is not None:
                            if re.match(inline_code_pattern2, text2):
                                # change formatting and add to paragraph object
                                text2 = text2.replace("`", "")
                                para_run = paragraph.add_run(text2)
                                font = para_run.font
                                font.name = 'consolas'
                                font.size = Pt(9)
                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                            # if italic text is in this section of text, format italic
                            elif re.findall(italic_text_pattern, text2):
                                italic_text_split = re.split(italic_text_pattern, text2)
                                for text3 in italic_text_split:
                                    if text3 is not None:
                                        if re.match(italic_text_pattern2, text3):
                                            # format italic and add to paragraph object
                                            text3 = text3.replace("*", "")
                                            para_run = paragraph.add_run(text3)
                                            para_run.italic = True
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # or if italic text is in the line, format italic
                elif re.findall(italic_text_pattern, text_):
                    italic_text = re.split(italic_text_pattern, text_)
                    for text2 in italic_text:
                        if text2 is not None:
                            if re.match(italic_text_pattern2, text2):
                                # format italic and add to paragraph object
                                text2 = text2.replace("*", "")
                                para_run = paragraph.add_run(text2)
                                para_run.italic = True
                            else:
                                paragraph.add_run(text2)
                else:
                    paragraph.add_run(text_)

    # convert bold and italic text
    # find all instances of bold and italic text in the line
    elif re.findall(italic_bold_text_pattern, line):
        italic_bold_text = re.split(italic_bold_text_pattern, line)
        # loop through the sections of text
        for text_ in italic_bold_text:
            if text_ is not None:
                # convert to bold and italic
                if re.match(italic_bold_text_pattern2, text_):
                    # format italic and add to paragraph object
                    text_ = text_.replace("***", "")
                    para_run = paragraph.add_run(text_)
                    para_run.italic = True
                    para_run.bold = True
                # or if bold text is in this section of text, format bold
                elif re.findall(bold_text_pattern, text_):
                    bold_text_split = re.split(bold_text_pattern, text_)
                    for text2 in bold_text_split:
                        if text2 is not None:
                            if re.match(bold_text_pattern2, text2):
                                # format bold and add to paragraph object
                                if '***' not in text2:
                                    text2 = text2.replace("**", "")
                                    para_run = paragraph.add_run(text2)
                                    para_run.bold = True
                            # or if italic text is in this section of text, format italic
                            elif re.findall(italic_text_pattern, text2):
                                italic_text_split = re.split(italic_text_pattern, text2)
                                for text3 in italic_text_split:
                                    if text3 is not None:
                                        if re.match(italic_text_pattern2, text3):
                                            # format italic and add to paragraph object
                                            if '***' not in text3 and '**' not in text3:
                                                text3 = text3.replace("*", "")
                                                para_run = paragraph.add_run(text3)
                                                para_run.italic = True
                                        # or if inline code is also in this section of text, format inline code
                                        elif re.findall(inline_code_pattern, text3):
                                            inline_code_split = re.split(inline_code_pattern, text3)
                                            for text4 in inline_code_split:
                                                if text4 is not None:
                                                    if re.match(inline_code_pattern2, text4):
                                                        # change formatting and add to paragraph object
                                                        text4 = text4.replace("`", "")
                                                        para_run = paragraph.add_run(text4)
                                                        font = para_run.font
                                                        font.name = 'consolas'
                                                        font.size = Pt(9)
                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                    else:
                                                        paragraph.add_run(text4)
                                        else:
                                            paragraph.add_run(text3)
                            # or if inline code is in this section of text, format inline code
                            elif re.findall(inline_code_pattern, text2):
                                inline_code_split = re.split(inline_code_pattern, text2)
                                for text3 in inline_code_split:
                                    if text3 is not None:
                                        if re.match(inline_code_pattern2, text3):
                                            # change formatting and add to paragraph object
                                            text3 = text3.replace("`", "")
                                            para_run = paragraph.add_run(text3)
                                            font = para_run.font
                                            font.name = 'consolas'
                                            font.size = Pt(9)
                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # or if italic text in this section of text, format italic
                elif re.findall(italic_text_pattern, text_):
                    italic_text_split = re.split(italic_text_pattern, text_)
                    for text2 in italic_text_split:
                        if text2 is not None:
                            if re.match(italic_text_pattern2, text2):
                                # format italic and add to docx para
                                if '***' not in text2 and '**' not in text2:
                                    text2 = text2.replace("*", "")
                                    para_run = paragraph.add_run(text2)
                                    para_run.italic = True
                            # or if inline code is in this section of text, format inline code
                            elif re.findall(inline_code_pattern, text2):
                                inline_code_split = re.split(inline_code_pattern, text2)
                                for text3 in inline_code_split:
                                    if text3 is not None:
                                        if re.match(inline_code_pattern2, text3):
                                            # change formatting and add to paragraph object
                                            text3 = text3.replace("`", "")
                                            para_run = paragraph.add_run(text3)
                                            font = para_run.font
                                            font.name = 'consolas'
                                            font.size = Pt(9)
                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # or if inline code is also in this section of text, format inline code
                elif re.findall(inline_code_pattern, text_):
                    inline_code_split = re.split(inline_code_pattern, text_)
                    for text2 in inline_code_split:
                        if text2 is not None:
                            if re.match(inline_code_pattern2, text2):
                                # change formatting and add to paragraph object
                                text2 = text2.replace("`", "")
                                para_run = paragraph.add_run(text2)
                                font = para_run.font
                                font.name = 'consolas'
                                font.size = Pt(9)
                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                            # or if italic text is in this section of text, format italic
                            elif re.findall(italic_text_pattern, text2):
                                italic_text_split = re.split(italic_text_pattern, text2)
                                for text3 in italic_text_split:
                                    if text3 is not None:
                                        if re.match(italic_text_pattern2, text3):
                                            # format italic and add to paragraph object
                                            if '***' not in text3 and '**' not in text3:
                                                text3 = text3.replace("*", "")
                                                para_run = paragraph.add_run(text3)
                                                para_run.italic = True
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                else:
                    paragraph.add_run(text_)

    # convert bold text
    # find any bold text in the markdown
    elif re.findall(bold_text_pattern, line):
        # split the line by the bold text
        bold_text_split = re.split(bold_text_pattern, line)
        for text_ in bold_text_split:
            if text_ is not None:
                # format bold and add to paragraph object
                if re.match(bold_text_pattern2, text_):
                    if '***' not in text_:
                        text_ = text_.replace("**", "")
                        para_run = paragraph.add_run(text_)
                        para_run.bold = True
                # or if italic text is in this section of text, format italic
                elif re.findall(italic_text_pattern, text_):
                    italic_text_split = re.split(italic_text_pattern, text_)
                    for text2 in italic_text_split:
                        if text2 is not None:
                            if re.match(italic_text_pattern2, text2):
                                # format italic and add to paragraph object
                                if '***' not in text2 and '**' not in text2:
                                    text2 = text2.replace("*", "")
                                    para_run = paragraph.add_run(text2)
                                    para_run.italic = True
                            # if inline code is in this section of text, format inline code
                            elif re.findall(inline_code_pattern, text2):
                                inline_code_split = re.split(inline_code_pattern, text2)
                                for text3 in inline_code_split:
                                    if text3 is not None:
                                        if re.match(inline_code_pattern2, text3):
                                            # change formatting and add to paragraph object
                                            text3 = text3.replace("`", "")
                                            para_run = paragraph.add_run(text3)
                                            font = para_run.font
                                            font.name = 'consolas'
                                            font.size = Pt(9)
                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        else:
                                            paragraph.add_run(text3)
                            else:
                                paragraph.add_run(text2)
                # if inline code is in this section of text, format inline code
                elif re.findall(inline_code_pattern, text_):
                    inline_code_split = re.split(inline_code_pattern, text_)
                    for text2 in inline_code_split:
                        if text2 is not None:
                            if re.match(inline_code_pattern2, text2):
                                text2 = text2.replace("`", "")
                                para_run = paragraph.add_run(text2)
                                font = para_run.font
                                font.name = 'consolas'
                                font.size = Pt(9)
                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                            else:
                                paragraph.add_run(text2)
                else:
                    paragraph.add_run(text_)

    # convert inline code
    # change the font of inline code to Consolas and add light grey highlighting
    elif re.findall(inline_code_pattern, line):
        inline_code_split = re.split(inline_code_pattern, line)
        for text_ in inline_code_split:
            if text_ is not None:
                # change formatting and add to paragraph object
                if re.match(inline_code_pattern2, text_):
                    text_ = text_.replace("`", "")
                    para_run = paragraph.add_run(text_)
                    font = para_run.font
                    font.name = 'consolas'
                    font.size = Pt(9)
                    font.highlight_color = WD_COLOR_INDEX.GRAY_25
                # if italic text is in this section of text, format italic
                elif re.findall(italic_text_pattern, text_):
                    italic_text_split = re.split(italic_text_pattern, text_)
                    for text2 in italic_text_split:
                        if text2 is not None:
                            if re.match(italic_text_pattern2, text2):
                                # format italic and add to paragraph object
                                if '***' not in text2 and '**' not in text2:
                                    text2 = text2.replace("*", "")
                                    para_run = paragraph.add_run(text2)
                                    para_run.italic = True
                            else:
                                paragraph.add_run(text2)
                else:
                    paragraph.add_run(text_)

    # if italic text is in this section of text, format italic
    elif re.findall(italic_text_pattern, line):
        italic_text = re.split(italic_text_pattern, line)
        for text_ in italic_text:
            if text_ is not None:
                if re.match(italic_text_pattern2, text_):
                    # format italic and add to paragraph object
                    if '***' not in text_ and '**' not in text_:
                        text_ = text_.replace("*", "")
                        para_run = paragraph.add_run(text_)
                        para_run.italic = True
                # if inline code is is in this section of text, format inline code
                elif re.findall(inline_code_pattern, text_):
                    inline_code_split = re.split(inline_code_pattern, text_)
                    for text2 in inline_code_split:
                        if text2 is not None:
                            if re.match(inline_code_pattern2, text2):
                                # change formatting and add to paragraph object
                                text2 = text2.replace("`", "")
                                para_run = paragraph.add_run(text2)
                                font = para_run.font
                                font.name = 'consolas'
                                font.size = Pt(9)
                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                            else:
                                paragraph.add_run(text2)
                else:
                    paragraph.add_run(text_)
    else:
        paragraph.add_run(line)


def markdown_to_docx():
    """The main function to convert markdown to docx.

    It loops through each line of the markdown file searching for markdown elements, converting them to their docx
    equivalent and adding them to the docx document object.
    """

    # declare folder path
    url_path = input('Enter the path to the folder containing the markdown files: ')
    # check that the folder exists; quit if it doesn't
    if not os.path.isdir(url_path):
        input('\nThat folder doesn\'t exist! Press any key to exit.')
        quit()

    # create output folder and declare output folder path
    if not os.path.isdir(url_path + '\\' + 'Output docx'):
        os.mkdir(url_path + '\\' + 'Output docx')
        output_url_path = url_path + '\\' + 'Output docx/'
    else:
        output_url_path = url_path + '\\' + 'Output docx/'

    # set file count to 0
    file_count = 0
    # loop through files in url path
    for file in os.listdir(url_path):
        if file.lower().endswith('.md'):
            # confirmation
            print('Converting ' + file)
            file_count += 1
            # file name without extension
            file_no_ext = file.split('.md')[0]
            # read the file
            f = open(url_path+'\\'+file, 'r')
            html_markdown = f.read()
            # create a docx document object
            document = Document()
            # set paragraph spacing
            paragraph_format = document.styles['Normal'].paragraph_format
            paragraph_format.space_after = Pt(1)
            # split the markdown file into lines
            lines = html_markdown.split('\n')

            # make the first line the title heading (if the first line is not a head1)
            if '#' not in lines[0][:3]:
                document.add_heading(lines[0].replace('**',''), level=0)
                del lines[0]
            # loop through each line of markdown
            for line in lines:
                if line:
                    # convert heading 1
                    if line[:7].count('#') == 1:
                        line = line.replace('# ', '')
                        line = line.replace('#', '')
                        document.add_heading(line, level=1) # add to document object
                    # convert heading 2
                    elif line[:7].count('#') == 2:
                        line = line.replace('## ', '')
                        line = line.replace('##', '')
                        document.add_heading(line, level=2) # add to document object
                    # convert heading 3
                    elif line[:7].count('#') == 3:
                        line = line.replace('### ', '')
                        line = line.replace('###', '')
                        document.add_heading(line, level=3) # add to document object
                    # convert heading 4
                    elif line[:7].count('#') == 4:
                        line = line.replace('#### ', '')
                        line = line.replace('####', '')
                        document.add_heading(line, level=4) # add to document object
                    # convert heading 5
                    elif line[:7].count('#') == 5:
                        line = line.replace('##### ', '')
                        line = line.replace('#####', '')
                        document.add_heading(line, level=5) # add to document object
                    # convert heading 6
                    elif line[:7].count('#') == 6:
                        line = line.replace('###### ', '')
                        line = line.replace('######', '')
                        document.add_heading(line, level=6) # add to document object
                    # convert heading 7
                    elif line[:7].count('#') == 7:
                        line = line.replace('####### ', '')
                        line = line.replace('#######', '')
                        document.add_heading(line, level=7) # add to document object
                    # convert heading 8
                    elif line[:8].count('#') == 8:
                        line = line.replace('######## ', '')
                        line = line.replace('########', '')
                        document.add_heading(line, level=8) # add to document object

                    # add in horizontal rules
                    elif line[:3] == '---':
                        # add horizontal rule to document object
                        document.add_paragraph("_____________________________________________")

                    # indent block quotes
                    elif line[:2].count('>') ==1:
                        line = line.replace('>', '')
                        line = line.replace('> ', '')
                        paragraph = document.add_paragraph() # add paragraph object
                        paragraph.paragraph_format.left_indent = Inches(0.25) # indent para
                        # convert any formatting/inline code in the block quote and add to docx para
                        format_para(line, paragraph)
                    elif line[:3].count('>') ==2:
                        line = line.replace('>>', '')
                        line = line.replace('>> ', '')
                        line = line.replace('> >', '')
                        line = line.replace('> > ', '')
                        paragraph = document.add_paragraph() # add paragraph object
                        paragraph.paragraph_format.left_indent = Inches(0.5) # indent para
                        # convert any formatting/inline code in the block quote
                        format_para(line, paragraph)

                    # unordered lists...
                    # level 1
                    elif str(line[:2]) == '- ' or str(line[:2]) == '* ' or str(line[:2]) == '+ ':
                        line = line[2:]
                        # add para styled as level 1 list to document object
                        paragraph = document.add_paragraph(style='List Bullet')
                        # convert any formatting/inline code in the list item and add to paragraph object
                        format_para(line, paragraph)
                    # level 2
                    elif str(line[:4]) == '  - ' or str(line[:4]) == '  * ' or str(line[:4]) == '  + ':
                        line = line[4:]
                        # add para styled as level 2 list to document object
                        paragraph = document.add_paragraph(style='List Bullet 2')
                        # convert any formatting/inline code in the list item and add to paragraph object
                        format_para(line, paragraph)
                    # level 3
                    elif str(line[:6]) == '    - ' or str(line[:6]) == '    * ' or str(line[:6]) == '    + ':
                        line = line[6:]
                        # add para styled as level 3 list to document object
                        paragraph = document.add_paragraph(style='List Bullet 3')
                        # convert with any formatting/inline code in the list item
                        format_para(line, paragraph)

                    # ordered lists...
                    # level 1 numbered list
                    elif line[0].isdigit():
                        if line[1] == '.':
                            line = line[3:]
                            # add para styled as list number to document object
                            paragraph = document.add_paragraph(style='List Number')
                            # convert any formatting/inline code within the list item and add to paragraph object
                            format_para(line, paragraph)
                        elif line[1].isdigit() and line[2] == '.':
                            line = line[4:]
                            paragraph = document.add_paragraph(style='List Number')
                            # convert any formatting/inline code within the list item and add to paragraph object
                            format_para(line, paragraph)
                    # level 2 numbered list
                    elif re.findall(level2_ordered_list, line[:5]):
                        line = re.split(level2_ordered_list, line)[-1].strip()
                        # add para styled as level 2 numbered list to document object
                        paragraph = document.add_paragraph(style='List Number 2')
                        format_para(line, paragraph)
                    # level 3 ordered list
                    elif re.findall(level3_ordered_list, line[:7]):
                        line = re.split(level3_ordered_list, line)[-1].strip()
                        # add para styled as level 3 numbered list to document object
                        paragraph = document.add_paragraph(style='List Number 3')
                        format_para(line, paragraph)

                    # indent any block code, change font and size, and prefix each line of code with a vertical green line
                    elif line[:4] == '    ':
                        line = line[4:] # delete the 4 blank spaces from start of line
                        if line !='': # if the line is not empty
                            paragraph = document.add_paragraph() # add paragraph object
                            # prefix block code with a vertical green line
                            para_run_ = paragraph.add_run('|')
                            para_run_.font.color.rgb = RGBColor(124,252,0)
                            para_run_.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                            para_run = paragraph.add_run(line) # add line to para
                            paragraph.paragraph_format.left_indent = Inches(0.25) # indent para
                            # set para spacing
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = Pt(0)
                            font = para_run.font
                            font.name = 'Consolas' # change font
                            font.size = Pt(9) # change size
                            font.highlight_color = WD_COLOR_INDEX.GRAY_25 # highlight light grey

                    # convert any formatting/inline code/hyperlinks in the markdown line (where a new paragraph object is required)..

                    # convert hyperlinks
                    elif re.findall(hyperlink_pattern, line):
                        # add paragraph object
                        paragraph = document.add_paragraph()
                        hyperlink = re.split(hyperlink_pattern, line)
                        for text_ in hyperlink:
                            if text_ is not None:
                                if re.match(hyperlink_pattern, text_):
                                    # get the url
                                    url_ = re.split(hyperlink_url_pattern, text_)[-2]
                                    # get the hyperlink text
                                    name_ = ""
                                    if re.split(hyperlink_name_pattern, text_)[0] != '' and \
                                            re.split(hyperlink_name_pattern, text_)[0].startswith('['):
                                        name_ = re.split(hyperlink_name_pattern, text_)[0].strip('[')
                                    elif re.split(hyperlink_name_pattern, text_)[1] != '' and \
                                            re.split(hyperlink_name_pattern, text_)[1].startswith('['):
                                        name_ = re.split(hyperlink_name_pattern, text_)[1].strip('[')
                                    # add in the hyperlink and format text blue and underlined
                                    add_hyperlink(paragraph, url_, name_, '334BFF', True)

                                # or if bold, italic text is in this section of text, format bold, italic
                                elif re.findall(italic_bold_text_pattern, text_):
                                    italic_bold_text = re.split(italic_bold_text_pattern, text_)
                                    for text2 in italic_bold_text:
                                        if text2 is not None:
                                            if re.match(italic_bold_text_pattern2, text2):
                                                text2 = text2.replace("***", "")
                                                # convert to bold italic and add to paragraph object
                                                para_run = paragraph.add_run(text2)
                                                para_run.italic = True
                                                para_run.bold = True
                                            # or if bold text is in this section of text, format bold
                                            elif re.findall(bold_text_pattern, text2):
                                                bold_text_split = re.split(bold_text_pattern, text2)
                                                for text3 in bold_text_split:
                                                    if text3 is not None:
                                                        if re.match(bold_text_pattern2, text3):
                                                            # convert to bold and add to paragraph object
                                                            text3 = text3.replace("**", "")
                                                            para_run = paragraph.add_run(text3)
                                                            para_run.bold = True
                                                        # or if italic text is in this section of text, format italic
                                                        elif re.findall(italic_text_pattern, text3):
                                                            italic_text_split = re.split(italic_text_pattern, text3)
                                                            for text4 in italic_text_split:
                                                                if text4 is not None:
                                                                    if re.match(italic_text_pattern2, text4):
                                                                        # convert to italic and add to paragraph object
                                                                        text4 = text4.replace("*", "")
                                                                        para_run = paragraph.add_run(text4)
                                                                        para_run.italic = True
                                                                    # or if inline code is in this section of text, format inline code
                                                                    elif re.findall(inline_code_pattern, text4):
                                                                        inline_code_split = re.split(inline_code_pattern, text4)
                                                                        for text5 in inline_code_split:
                                                                            if text5 is not None:
                                                                                if re.match(inline_code_pattern2, text5):
                                                                                    # change formatting and add to paragraph object
                                                                                    text5 = text5.replace("`", "")
                                                                                    para_run = paragraph.add_run(text5)
                                                                                    font = para_run.font
                                                                                    font.name = 'consolas'
                                                                                    font.size = Pt(9)
                                                                                    font.highlight_color = \
                                                                                        WD_COLOR_INDEX.GRAY_25
                                                                                else:
                                                                                    paragraph.add_run(text5)
                                                                    else:
                                                                        paragraph.add_run(text4)
                                                        # or if inline code is in this section of text, format inline code
                                                        elif re.findall(inline_code_pattern, text3):
                                                            inline_code_split = re.split(inline_code_pattern, text3)
                                                            for text4 in inline_code_split:
                                                                if text4 is not None:
                                                                    if re.match(inline_code_pattern2, text4):
                                                                        # change formatting and add to paragraph object
                                                                        text4 = text4.replace("`", "")
                                                                        para_run = paragraph.add_run(text4)
                                                                        font = para_run.font
                                                                        font.name = 'consolas'
                                                                        font.size = Pt(9)
                                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                                    else:
                                                                        paragraph.add_run(text4)
                                                        else:
                                                            paragraph.add_run(text3)

                                            # or if italic text is in this section of text, format italic
                                            elif re.findall(italic_text_pattern, text2):
                                                italic_text_split = re.split(italic_text_pattern, text2)
                                                for text3 in italic_text_split:
                                                    if text3 is not None:
                                                        if re.match(italic_text_pattern2, text3):
                                                            # convert to italic and add to paragraph object
                                                            text3 = text3.replace("*", "")
                                                            para_run = paragraph.add_run(text3)
                                                            para_run.italic = True
                                                        # or if inline code is in this section of text, format inline code
                                                        elif re.findall(inline_code_pattern, text3):
                                                            inline_code_split = re.split(inline_code_pattern, text3)
                                                            for text4 in inline_code_split:
                                                                if text4 is not None:
                                                                    if re.match(inline_code_pattern2, text4):
                                                                        # change formatting and add to paragraph object
                                                                        text4 = text4.replace("`", "")
                                                                        para_run = paragraph.add_run(text4)
                                                                        font = para_run.font
                                                                        font.name = 'consolas'
                                                                        font.size = Pt(9)
                                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                                    else:
                                                                        paragraph.add_run(text4)
                                                        else:
                                                            paragraph.add_run(text3)
                                            # if inline code is in this section of text, format inline code
                                            elif re.findall(inline_code_pattern, text2):
                                                inline_code_split = re.split(inline_code_pattern, text2)
                                                for text3 in inline_code_split:
                                                    if text3 is not None:
                                                        if re.match(inline_code_pattern2, text3):
                                                            # change formatting and add to paragraph object
                                                            text3 = text3.replace("`", "")
                                                            para_run = paragraph.add_run(text3)
                                                            font = para_run.font
                                                            font.name = 'consolas'
                                                            font.size = Pt(9)
                                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)

                                # or if bold text is in the line, format bold
                                elif re.findall(bold_text_pattern, text_):
                                    bold_text_split = re.split(bold_text_pattern, text_)
                                    for text2 in bold_text_split:
                                        if text2 is not None:
                                            if re.match(bold_text_pattern2, text2):
                                                # format bold and add to paragraph object
                                                text2 = text2.replace("**", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.bold = True
                                            # or if italic text is in this section of text, format italic
                                            elif re.findall(italic_text_pattern, text2):
                                                italic_text_split = re.split(italic_text_pattern, text2)
                                                for text3 in italic_text_split:
                                                    if text3 is not None:
                                                        if re.match(italic_text_pattern2, text3):
                                                            # format italic and add to paragraph object
                                                            text3 = text3.replace("*", "")
                                                            para_run = paragraph.add_run(text3)
                                                            para_run.italic = True
                                                        # if inline code is in this section of text, format italic
                                                        elif re.findall(inline_code_pattern, text3):
                                                            inline_code_split = re.split(inline_code_pattern, text3)
                                                            for text4 in inline_code_split:
                                                                if text4 is not None:
                                                                    if re.match(inline_code_pattern2,text4):
                                                                        # change formatting and add to paragraph object
                                                                        text4 = text4.replace("`", "")
                                                                        para_run = paragraph.add_run(text4)
                                                                        font = para_run.font
                                                                        font.name = 'consolas'
                                                                        font.size = Pt(9)
                                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                                    else:
                                                                        paragraph.add_run(text4)
                                                        else:
                                                            paragraph.add_run(text3)

                                            # or if inline code is in this section of text, format inline code
                                            elif re.findall(inline_code_pattern, text2):
                                                inline_code_split = re.split(inline_code_pattern, text2)
                                                for text3 in inline_code_split:
                                                    if text3 is not None:
                                                        if re.match(inline_code_pattern2, text3):
                                                            # change formatting and add to to paragraph object
                                                            text3 = text3.replace("`", "")
                                                            para_run = paragraph.add_run(text3)
                                                            font = para_run.font
                                                            font.name = 'consolas'
                                                            font.size = Pt(9)
                                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)

                                # inline code
                                # change the font of inline code in the line to Consolas and highlight in light grey
                                elif re.findall(inline_code_pattern, text_):
                                    inline_code_split = re.split(inline_code_pattern, text_)
                                    for text2 in inline_code_split:
                                        if text2 is not None:
                                            if re.match(inline_code_pattern2, text2):
                                                # change formatting and add to paragraph object
                                                text2 = text2.replace("`", "")
                                                para_run = paragraph.add_run(text2)
                                                font = para_run.font
                                                font.name = 'consolas'
                                                font.size = Pt(9)
                                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                            # if italic text is in this section of text, format italic
                                            elif re.findall(italic_text_pattern, text2):
                                                italic_text_split = re.split(italic_text_pattern, text2)
                                                for text3 in italic_text_split:
                                                    if text3 is not None:
                                                        if re.match(italic_text_pattern2, text3):
                                                            # format italic and add to paragraph object
                                                            text3 = text3.replace("*", "")
                                                            para_run = paragraph.add_run(text3)
                                                            para_run.italic = True
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)

                                # or if italic text is in the line, format italic
                                elif re.findall(italic_text_pattern, text_):
                                    italic_text = re.split(italic_text_pattern, text_)
                                    for text2 in italic_text:
                                        if text2 is not None:
                                            if re.match(italic_text_pattern2, text2):
                                                # format italic and add to paragraph object
                                                text2 = text2.replace("*", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.italic = True
                                            else:
                                                paragraph.add_run(text2)
                                else:
                                    paragraph.add_run(text_)
                    #---------------------------------------------------------------------------------------------------
                    # if bold, italic text is in the line, format bold, italic
                    elif re.findall(italic_bold_text_pattern, line):
                        paragraph = document.add_paragraph() # add para to paragraph object
                        italic_bold_text = re.split(italic_bold_text_pattern, line)
                        for text_ in italic_bold_text:
                            if text_ is not None:
                                if re.match(italic_bold_text_pattern2, text_):
                                    text_ = text_.replace("***", "")
                                    # convert to bold italic and add to paragraph object
                                    para_run = paragraph.add_run(text_)
                                    para_run.italic = True
                                    para_run.bold = True
                                # or if bold text is in this section of text, format bold
                                elif re.findall(bold_text_pattern, text_):
                                    bold_text_split = re.split(bold_text_pattern, text_)
                                    for text2 in bold_text_split:
                                        if text2 is not None:
                                            if re.match(bold_text_pattern2, text2):
                                                # convert to bold and add to paragraph object
                                                text2 = text2.replace("**", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.bold = True
                                            # or if italic text is in this section of text, format italic
                                            elif re.findall(italic_text_pattern, text2):
                                                italic_text_split = re.split(italic_text_pattern, text2)
                                                for text3 in italic_text_split:
                                                    if text3 is not None:
                                                        if re.match(italic_text_pattern2, text3):
                                                            # convert to italic and add to paragraph object
                                                            text3 = text3.replace("*", "")
                                                            para_run = paragraph.add_run(text3)
                                                            para_run.italic = True
                                                        # or if inline code is in this section of text, format inline code
                                                        elif re.findall(inline_code_pattern, text3):
                                                            inline_code_split = re.split(inline_code_pattern, text3)
                                                            for text4 in inline_code_split:
                                                                if text4 is not None:
                                                                    if re.match(inline_code_pattern2, text4):
                                                                        # change formatting and add to paragraph object
                                                                        text4 = text4.replace("`", "")
                                                                        para_run = paragraph.add_run(text4)
                                                                        font = para_run.font
                                                                        font.name = 'consolas'
                                                                        font.size = Pt(9)
                                                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                                    else:
                                                                        paragraph.add_run(text4)
                                                        else:
                                                            paragraph.add_run(text3)
                                            # or if inline code is in this section of text, format inline code
                                            elif re.findall(inline_code_pattern, text2):
                                                inline_code_split = re.split(inline_code_pattern, text2)
                                                for text3 in inline_code_split:
                                                    if text3 is not None:
                                                        if re.match(inline_code_pattern2, text3):
                                                            # change formatting and add to paragraph object
                                                            text3 = text3.replace("`", "")
                                                            para_run = paragraph.add_run(text3)
                                                            font = para_run.font
                                                            font.name = 'consolas'
                                                            font.size = Pt(9)
                                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)

                                # or if italic text is in this section of text, format italic
                                elif re.findall(italic_text_pattern, text_):
                                    italic_text_split = re.split(italic_text_pattern, text_)
                                    for text2 in italic_text_split:
                                        if text2 is not None:
                                            if re.match(italic_text_pattern2, text2):
                                                # convert to italic and add to paragraph object
                                                text2 = text2.replace("*", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.italic = True
                                            # or if inline code is in this section of text, format inline code
                                            elif re.findall(inline_code_pattern, text2):
                                                inline_code_split = re.split(inline_code_pattern, text2)
                                                for text3 in inline_code_split:
                                                    if text3 is not None:
                                                        if re.match(inline_code_pattern2, text3):
                                                            # change formatting and add to paragraph object
                                                            text3 = text3.replace("`", "")
                                                            para_run = paragraph.add_run(text3)
                                                            font = para_run.font
                                                            font.name = 'consolas'
                                                            font.size = Pt(9)
                                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)
                                # if inline code is in this section of text, format inline code
                                elif re.findall(inline_code_pattern, text_):
                                    inline_code_split = re.split(inline_code_pattern, text_)
                                    for text2 in inline_code_split:
                                        if text2 is not None:
                                            if re.match(inline_code_pattern2, text2):
                                                # change formatting and add to paragraph object
                                                text2 = text2.replace("`", "")
                                                para_run = paragraph.add_run(text2)
                                                font = para_run.font
                                                font.name = 'consolas'
                                                font.size = Pt(9)
                                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                            else:
                                                paragraph.add_run(text2)
                                else:
                                    paragraph.add_run(text_)
                    #---------------------------------------------------------------------------------------------------
                    # or if bold text is in the line, format bold
                    elif re.findall(bold_text_pattern, line):
                        # add para to docx document
                        paragraph = document.add_paragraph()
                        bold_text_split = re.split(bold_text_pattern, line)
                        for text_ in bold_text_split:
                            if text_ is not None:
                                if re.match(bold_text_pattern2, text_):
                                    # format bold and add to paragraph object
                                    text_ = text_.replace("**", "")
                                    para_run = paragraph.add_run(text_)
                                    para_run.bold = True
                                # or if italic text is in this section of text, format italic
                                elif re.findall(italic_text_pattern, text_):
                                    italic_text_split = re.split(italic_text_pattern, text_)
                                    for text2 in italic_text_split:
                                        if text2 is not None:
                                            if re.match(italic_text_pattern2, text2):
                                                # format italic and add to paragraph object
                                                text2 = text2.replace("*", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.italic = True
                                            # if inline code is in this section of text, format italic
                                            elif re.findall(inline_code_pattern, text2):
                                                inline_code_split = re.split(inline_code_pattern, text2)
                                                for text3 in inline_code_split:
                                                    if text3 is not None:
                                                        if re.match(inline_code_pattern2, text3):
                                                            # change formatting and add to paragraph object
                                                            text3 = text3.replace("`", "")
                                                            para_run = paragraph.add_run(text3)
                                                            font = para_run.font
                                                            font.name = 'consolas'
                                                            font.size = Pt(9)
                                                            font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                                        else:
                                                            paragraph.add_run(text3)
                                            else:
                                                paragraph.add_run(text2)

                                # or if inline code is in this section of text, format inline code
                                elif re.findall(inline_code_pattern, text_):
                                    inline_code_split = re.split(inline_code_pattern, text_)
                                    for text2 in inline_code_split:
                                        if text2 is not None:
                                            if re.match(inline_code_pattern2, text2):
                                                # change formatting and add to to paragraph object
                                                text2  = text2 .replace("`", "")
                                                para_run = paragraph.add_run(text2)
                                                font = para_run.font
                                                font.name = 'consolas'
                                                font.size = Pt(9)
                                                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                            else:
                                                paragraph.add_run(text2)
                                else:
                                    paragraph.add_run(text_)
                    #---------------------------------------------------------------------------------------------------
                    # inline code
                    # change the font of inline code in the line to 9pt Consolas and highlighted in light grey
                    elif re.findall(inline_code_pattern, line):
                        paragraph = document.add_paragraph()
                        inline_code_split = re.split(inline_code_pattern, line)
                        for text_ in inline_code_split:
                            if text_ is not None:
                                if re.match(inline_code_pattern2, text_):
                                    # change formatting and add to paragraph object
                                    text_ = text_.replace("`", "")
                                    para_run = paragraph.add_run(text_)
                                    font = para_run.font
                                    font.name = 'consolas'
                                    font.size = Pt(9)
                                    font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                # if italic text is in this section of text, format italic
                                elif re.findall(italic_text_pattern, text_):
                                    italic_text_split = re.split(italic_text_pattern, text_)
                                    for text2 in italic_text_split:
                                        if text2 is not None:
                                            if re.match(italic_text_pattern2, text2):
                                                # format italic and add to paragraph object
                                                text2 = text2.replace("*", "")
                                                para_run = paragraph.add_run(text2)
                                                para_run.italic = True
                                            else:
                                                paragraph.add_run(text2)
                                else:
                                    paragraph.add_run(text_)
                    #---------------------------------------------------------------------------------------------------
                    # or if italic text is in the line, format italic
                    elif re.findall(italic_text_pattern, line):
                        paragraph = document.add_paragraph()
                        italic_text = re.split(italic_text_pattern, line)
                        for text_ in italic_text:
                            if text_ is not None:
                                if re.match(italic_text_pattern2, text_):
                                    # format italic and add to paragraph object
                                    text_ = text_.replace("*", "")
                                    para_run = paragraph.add_run(text_)
                                    para_run.italic = True

                                else:
                                    paragraph.add_run(text_)
                    else:
                        paragraph = document.add_paragraph()
                        paragraph.add_run(line)
                else:
                    document.add_paragraph()

            # save document object to docx file
            document.save(output_url_path + file_no_ext+'.docx')

    # get the end time of the script
    end_time = time.time()
    end_msg_txt = '\n' + endmessage.getmessagetext(file_count) + scripttime.runningtime(start_time,end_time)
    # show confirmation message and exit prompt
    input(end_msg_txt + '\n\nThe converted docx files are saved in an "Output docx" folder'
                        ': '+output_url_path+ '\n\nPress any key to exit………')

# call the main markdown_to_docx function
if __name__ == '__main__':
    markdown_to_docx()
